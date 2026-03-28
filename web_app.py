#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FormulaAI Web - AI智能文档排版工具 (Web版)

基于Streamlit构建的Web GUI，复用核心模块。
支持：文档排版、模板管理、页眉页脚配置、API配置持久化
"""

import streamlit as st
import tempfile
import io
import os
import sys
import json
from datetime import datetime

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入核心模块
from src.core.ai_connector import AIConnector
from src.core.format_manager import FormatManager
from src.core.doc_processor import DocProcessor
from src.core.structure_analyzer import StructureAnalyzer
from src.core.header_footer_config import HeaderFooterConfig
from src.utils.config_manager import ConfigManager
from src.utils.logger import app_logger


# ========================
# 页面配置
# ========================
st.set_page_config(
    page_title="FormulaAI Web",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========================
# 样式设置
# ========================
st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: bold;
        text-align: center;
        margin-bottom: 1rem;
    }
    .success-box {
        padding: 1rem;
        background-color: #e8f5e9;
        border-radius: 0.5rem;
        border-left: 5px solid #4caf50;
    }
    .error-box {
        padding: 1rem;
        background-color: #ffebee;
        border-radius: 0.5rem;
        border-left: 5px solid #f44336;
    }
    .info-box {
        padding: 1rem;
        background-color: #e3f2fd;
        border-radius: 0.5rem;
        border-left: 5px solid #2196f3;
    }
    .template-card {
        padding: 1rem;
        border: 1px solid #e0e0e0;
        border-radius: 0.5rem;
        margin-bottom: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)


# ========================
# 会话状态初始化
# ========================
def init_session_state():
    """初始化会话状态"""
    # API配置
    if 'api_url' not in st.session_state:
        st.session_state.api_url = ""
    if 'api_key' not in st.session_state:
        st.session_state.api_key = ""
    if 'model' not in st.session_state:
        st.session_state.model = "deepseek-chat"

    # 文档处理
    if 'uploaded_file' not in st.session_state:
        st.session_state.uploaded_file = None
    if 'output_bytes' not in st.session_state:
        st.session_state.output_bytes = None
    if 'logs' not in st.session_state:
        st.session_state.logs = []

    # 模板管理
    if 'editing_template' not in st.session_state:
        st.session_state.editing_template = None
    if 'new_template_mode' not in st.session_state:
        st.session_state.new_template_mode = False

    # 页眉页脚配置
    if 'header_footer_config' not in st.session_state:
        st.session_state.header_footer_config = HeaderFooterConfig().to_dict()


init_session_state()

# 加载持久化配置
config_manager = ConfigManager()


# ========================
# 辅助函数
# ========================
def add_log(message: str, level: str = "INFO"):
    """添加日志"""
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.logs.append(f"[{timestamp}] [{level}] {message}")


def load_api_config():
    """加载API配置"""
    api_config = config_manager.get_api_config()
    if api_config:
        st.session_state.api_url = api_config.get('api_url', '')
        st.session_state.api_key = api_config.get('api_key', '')
        st.session_state.model = api_config.get('model', 'deepseek-chat')


def save_api_config():
    """保存API配置"""
    config_manager.save_api_config({
        'api_url': st.session_state.api_url,
        'api_key': st.session_state.api_key,
        'model': st.session_state.model
    })


def load_header_footer_config():
    """加载页眉页脚配置"""
    app_config = config_manager.get_app_config()
    if 'header_footer_config' in app_config:
        st.session_state.header_footer_config = app_config['header_footer_config']


def save_header_footer_config():
    """保存页眉页脚配置"""
    app_config = config_manager.get_app_config()
    app_config['header_footer_config'] = st.session_state.header_footer_config
    config_manager.save_app_config(app_config)


# 初始化时加载配置
load_api_config()
load_header_footer_config()


# ========================
# 文档处理函数
# ========================
def process_document(uploaded_file, template_name: str, api_url: str, api_key: str, model: str, hf_config: dict):
    """处理文档排版"""
    add_log("开始处理文档...")

    # 1. 保存上传文件到临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name

    add_log(f"临时文件已创建: {tmp_path}")

    try:
        # 2. 读取文档
        doc = Document(tmp_path)
        paragraphs_text = [p.text for p in doc.paragraphs]
        add_log(f"文档读取成功，共 {len(paragraphs_text)} 个段落")

        # 3. 获取模板
        format_manager = FormatManager()
        template = format_manager.get_template(template_name)
        if not template:
            raise ValueError(f"模板 '{template_name}' 不存在")
        add_log(f"已加载模板: {template_name}")

        # 4. 调用AI API
        api_config = {
            "api_url": api_url,
            "api_key": api_key,
            "model": model,
            "timeout": 300
        }
        ai_connector = AIConnector(api_config)

        # 验证配置
        valid, error_msg = ai_connector.validate_config()
        if not valid:
            raise ValueError(f"API配置无效: {error_msg}")

        add_log("正在调用AI API...")

        # 生成提示词
        prompt = ai_connector.generate_prompt(paragraphs_text, template.get('rules', {}))

        # 发送请求
        success, response = ai_connector.send_request(prompt)
        if not success:
            raise ValueError(f"AI请求失败: {response}")

        add_log("AI响应已接收")

        # 解析响应
        success, formatting_instructions = ai_connector.parse_response(response)
        if not success:
            raise ValueError(f"解析响应失败: {formatting_instructions}")

        add_log(f"排版指令已生成，共 {len(formatting_instructions.get('elements', []))} 个元素")

        # 5. 应用排版格式
        output_doc = apply_formatting(doc, formatting_instructions, template, hf_config)

        # 6. 保存到字节流
        output_bytes = io.BytesIO()
        output_doc.save(output_bytes)
        output_bytes.seek(0)

        add_log("文档排版完成！")

        return output_bytes.getvalue()

    finally:
        # 清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def apply_formatting(original_doc, formatting_instructions, template, hf_config: dict):
    """应用排版格式到文档"""
    new_doc = Document()

    # 字号映射
    font_size_mapping = {
        "初号": Pt(42), "小初": Pt(36), "一号": Pt(26), "小一": Pt(24),
        "二号": Pt(22), "小二": Pt(18), "三号": Pt(16), "小三": Pt(15),
        "四号": Pt(14), "小四": Pt(12), "五号": Pt(10.5), "小五": Pt(9),
        "六号": Pt(7.5), "小六": Pt(6.5)
    }

    # 对齐映射
    alignment_mapping = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }

    elements = formatting_instructions.get('elements', [])

    for element in elements:
        content = element.get('content', '')
        format_info = element.get('format', {})

        # 创建段落
        para = new_doc.add_paragraph()
        run = para.add_run(content)

        # 应用字体
        font_name = format_info.get('font', '宋体')
        run.font.name = font_name

        # 尝试设置中文字体
        try:
            from docx.oxml.ns import qn
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except:
            pass

        # 应用字号
        font_size = format_info.get('size', '小四')
        if font_size in font_size_mapping:
            run.font.size = font_size_mapping[font_size]

        # 应用粗体
        run.font.bold = format_info.get('bold', False)

        # 应用斜体
        run.font.italic = format_info.get('italic', False)

        # 应用对齐
        alignment = format_info.get('alignment', 'left')
        para.alignment = alignment_mapping.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

    # TODO: 应用页眉页脚（简化版暂不支持）

    return new_doc


# ========================
# 页面：文档排版
# ========================
def page_document_format():
    """文档排版页面"""
    st.markdown('<p class="main-header">📄 文档排版</p>', unsafe_allow_html=True)

    # 检查API配置
    if not st.session_state.api_url or not st.session_state.api_key:
        st.warning("⚠️ 请先在「API配置」页面设置API参数")
        st.stop()

    col1, col2 = st.columns([2, 1])

    with col1:
        st.subheader("📤 上传文档")
        uploaded_file = st.file_uploader(
            "选择Word文档",
            type=['docx'],
            help="支持.docx格式的Word文档"
        )

        if uploaded_file:
            st.success(f"已选择: {uploaded_file.name} ({uploaded_file.size / 1024:.1f} KB)")
            st.session_state.uploaded_file = uploaded_file

    with col2:
        st.subheader("📝 模板选择")
        format_manager = FormatManager()
        template_names = format_manager.get_template_names()

        if template_names:
            selected_template = st.selectbox(
                "选择排版模板",
                template_names,
                index=0
            )

            # 显示模板预览
            if selected_template:
                template = format_manager.get_template(selected_template)
                if template:
                    with st.expander("模板预览"):
                        st.json(template.get('rules', {}))
        else:
            st.warning("没有可用的模板，请先创建模板")
            selected_template = None

    st.divider()

    # 页眉页脚设置（折叠）
    with st.expander("⚙️ 页眉页脚设置（可选）"):
        hf_config = st.session_state.header_footer_config

        c1, c2 = st.columns(2)
        with c1:
            enable_header = st.checkbox("启用页眉", value=hf_config.get('enable_header', False))
            header_content = st.text_input("页眉内容", value=hf_config.get('header_content', ''))
        with c2:
            enable_footer = st.checkbox("启用页脚", value=hf_config.get('enable_footer', False))
            footer_content = st.text_input("页脚内容", value=hf_config.get('footer_content', ''))

        # 更新配置
        st.session_state.header_footer_config.update({
            'enable_header': enable_header,
            'header_content': header_content,
            'enable_footer': enable_footer,
            'footer_content': footer_content
        })

    # 排版按钮
    col1, col2, col3 = st.columns([1, 1, 1])

    with col1:
        can_format = uploaded_file and selected_template and st.session_state.api_url and st.session_state.api_key
        start_button = st.button(
            "🚀 开始排版",
            type="primary",
            use_container_width=True,
            disabled=not can_format
        )

    with col2:
        if st.button("🗑️ 清除", use_container_width=True):
            st.session_state.uploaded_file = None
            st.session_state.output_bytes = None
            st.session_state.logs = []
            st.rerun()

    with col3:
        if st.session_state.output_bytes:
            st.download_button(
                label="📥 下载结果",
                data=st.session_state.output_bytes,
                file_name=f"formatted_{uploaded_file.name if uploaded_file else 'document.docx'}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    # 处理排版
    if start_button:
        st.session_state.logs = []

        with st.spinner("正在处理..."):
            progress_bar = st.progress(0)
            status_text = st.empty()

            try:
                status_text.text("正在处理文档...")
                progress_bar.progress(20)

                output_bytes = process_document(
                    uploaded_file,
                    selected_template,
                    st.session_state.api_url,
                    st.session_state.api_key,
                    st.session_state.model,
                    st.session_state.header_footer_config
                )

                progress_bar.progress(100)
                st.session_state.output_bytes = output_bytes

                st.success("✅ 排版完成！点击上方按钮下载结果。")

            except Exception as e:
                st.error(f"❌ 处理失败: {str(e)}")
                add_log(f"错误: {str(e)}", "ERROR")

    # 日志显示
    if st.session_state.logs:
        st.divider()
        st.subheader("📋 处理日志")
        with st.container(height=200):
            for log in st.session_state.logs:
                if "[ERROR]" in log:
                    st.error(log)
                elif "[WARNING]" in log:
                    st.warning(log)
                else:
                    st.text(log)


# ========================
# 页面：模板管理
# ========================
def page_template_management():
    """模板管理页面"""
    st.markdown('<p class="main-header">📝 模板管理</p>', unsafe_allow_html=True)

    format_manager = FormatManager()

    # 操作按钮
    col1, col2, col3 = st.columns([1, 1, 2])

    with col1:
        if st.button("➕ 新建模板", type="primary", use_container_width=True):
            st.session_state.new_template_mode = True
            st.session_state.editing_template = {
                'name': '',
                'description': '',
                'rules': {}
            }

    with col2:
        if st.button("🔄 刷新列表", use_container_width=True):
            format_manager.load_templates()
            st.rerun()

    st.divider()

    # 新建/编辑模板表单
    if st.session_state.new_template_mode or st.session_state.editing_template:
        template_editor(format_manager)
        return

    # 模板列表
    templates = format_manager.get_templates()

    if not templates:
        st.info("暂无模板，请点击「新建模板」创建")
        return

    for name, template in templates.items():
        with st.container():
            col1, col2, col3 = st.columns([3, 1, 1])

            with col1:
                st.markdown(f"**{name}**")
                if template.get('description'):
                    st.caption(template['description'])

            with col2:
                if st.button("✏️ 编辑", key=f"edit_{name}", use_container_width=True):
                    st.session_state.editing_template = template.copy()
                    st.session_state.new_template_mode = False
                    st.rerun()

            with col3:
                if st.button("🗑️ 删除", key=f"del_{name}", use_container_width=True):
                    if format_manager.delete_template(name):
                        st.success(f"已删除模板: {name}")
                        st.rerun()
                    else:
                        st.error("删除失败")

            # 显示模板规则
            with st.expander("查看规则"):
                rules = template.get('rules', {})
                for element_type, rule in rules.items():
                    st.markdown(f"**{element_type}**: {rule.get('font', '宋体')}, {rule.get('size', '小四')}")

            st.divider()


def template_editor(format_manager: FormatManager):
    """模板编辑器"""
    template = st.session_state.editing_template or {
        'name': '',
        'description': '',
        'rules': {}
    }

    is_new = st.session_state.new_template_mode

    st.subheader("🆕 新建模板" if is_new else f"✏️ 编辑模板: {template.get('name', '')}")

    # 基本信息
    name = st.text_input("模板名称", value=template.get('name', ''), disabled=not is_new)
    description = st.text_input("模板描述", value=template.get('description', ''))

    # 规则编辑
    st.subheader("格式规则")

    rules = template.get('rules', {})
    element_types = ['标题', '一级标题', '二级标题', '三级标题', '正文', '摘要', '关键词', '参考文献']

    # 字体选项
    fonts = ['宋体', '黑体', '楷体', '仿宋', 'Times New Roman', 'Arial']
    # 字号选项
    sizes = ['初号', '小初', '一号', '小一', '二号', '小二', '三号', '小三', '四号', '小四', '五号', '小五', '六号']
    # 对齐选项
    alignments = ['left', 'center', 'right', 'justify']
    alignment_names = {'left': '左对齐', 'center': '居中', 'right': '右对齐', 'justify': '两端对齐'}

    # 添加新规则
    with st.expander("➕ 添加新规则"):
        new_type = st.selectbox("元素类型", element_types, key="new_rule_type")

        col1, col2, col3 = st.columns(3)
        with col1:
            new_font = st.selectbox("字体", fonts, key="new_rule_font")
        with col2:
            new_size = st.selectbox("字号", sizes, key="new_rule_size")
        with col3:
            new_align = st.selectbox("对齐", alignments, format_func=lambda x: alignment_names[x], key="new_rule_align")

        new_bold = st.checkbox("粗体", key="new_rule_bold")

        if st.button("添加规则", key="add_rule_btn"):
            rules[new_type] = {
                'font': new_font,
                'size': new_size,
                'alignment': new_align,
                'bold': new_bold
            }
            template['rules'] = rules

    # 显示现有规则
    st.markdown("**当前规则:**")
    for element_type, rule in list(rules.items()):
        col1, col2, col3 = st.columns([3, 4, 1])
        with col1:
            st.markdown(f"**{element_type}**")
        with col2:
            st.text(f"{rule.get('font', '宋体')}, {rule.get('size', '小四')}, {alignment_names.get(rule.get('alignment', 'left'), '左对齐')}")
        with col3:
            if st.button("🗑️", key=f"del_rule_{element_type}"):
                del rules[element_type]
                template['rules'] = rules
                st.rerun()

    st.divider()

    # 保存按钮
    col1, col2 = st.columns(2)

    with col1:
        if st.button("💾 保存模板", type="primary", use_container_width=True):
            if not name:
                st.error("请输入模板名称")
            elif not rules:
                st.error("请至少添加一条规则")
            else:
                template['name'] = name
                template['description'] = description
                template['rules'] = rules

                if format_manager.save_template(name, template):
                    st.success(f"模板已保存: {name}")
                    st.session_state.editing_template = None
                    st.session_state.new_template_mode = False
                    st.rerun()
                else:
                    st.error("保存失败")

    with col2:
        if st.button("❌ 取消", use_container_width=True):
            st.session_state.editing_template = None
            st.session_state.new_template_mode = False
            st.rerun()


# ========================
# 页面：API配置
# ========================
def page_api_config():
    """API配置页面"""
    st.markdown('<p class="main-header">⚙️ API配置</p>', unsafe_allow_html=True)

    st.info("配置AI API参数，用于文档智能排版")

    # API URL
    api_url = st.text_input(
        "API URL (Base URL)",
        value=st.session_state.api_url,
        placeholder="https://api.deepseek.com",
        help="API基础地址，如: https://api.deepseek.com 或 https://api.openai.com"
    )

    # API Key
    api_key = st.text_input(
        "API Key",
        value=st.session_state.api_key,
        type="password",
        placeholder="sk-xxxxxxxxxxxxxxxx",
        help="您的API密钥"
    )

    # 模型选择
    common_models = [
        "deepseek-chat",
        "deepseek-coder",
        "gpt-3.5-turbo",
        "gpt-4",
        "gpt-4-turbo",
        "claude-3-opus",
        "claude-3-sonnet"
    ]

    model_option = st.radio(
        "模型选择方式",
        ["选择预设模型", "输入自定义模型"],
        horizontal=True
    )

    if model_option == "选择预设模型":
        model = st.selectbox(
            "模型",
            common_models,
            index=0 if st.session_state.model not in common_models else common_models.index(st.session_state.model)
        )
    else:
        model = st.text_input(
            "自定义模型名称",
            value=st.session_state.model,
            placeholder="model-name"
        )

    # 更新会话状态
    st.session_state.api_url = api_url
    st.session_state.api_key = api_key
    st.session_state.model = model

    st.divider()

    # 保存按钮
    col1, col2 = st.columns(2)

    with col1:
        if st.button("💾 保存配置", type="primary", use_container_width=True):
            save_api_config()
            st.success("配置已保存")

    with col2:
        if st.button("🔌 测试连接", use_container_width=True):
            if not api_url or not api_key:
                st.error("请填写API URL和Key")
            else:
                with st.spinner("测试中..."):
                    try:
                        api_config = {
                            "api_url": api_url,
                            "api_key": api_key,
                            "model": model
                        }
                        connector = AIConnector(api_config)
                        valid, msg = connector.validate_config()
                        if valid:
                            st.success(f"连接成功！模型: {model}")
                        else:
                            st.error(f"连接失败: {msg}")
                    except Exception as e:
                        st.error(f"测试失败: {str(e)}")

    # 显示当前配置
    st.divider()
    st.subheader("当前配置")

    api_config = config_manager.get_api_config()
    if api_config:
        st.json({
            "api_url": api_config.get('api_url', '(未设置)'),
            "model": api_config.get('model', '(未设置)'),
            "last_updated": api_config.get('last_updated', '(未知)')
        })
    else:
        st.text("暂无保存的配置")


# ========================
# 页面：从文本解析模板
# ========================
def page_text_parsing():
    """从文本解析模板页面"""
    st.markdown('<p class="main-header">🔮 从文本解析模板</p>', unsafe_allow_html=True)

    st.info("输入自然语言描述的格式要求，AI将自动生成排版模板")

    # 检查API配置
    if not st.session_state.api_url or not st.session_state.api_key:
        st.warning("⚠️ 请先在「API配置」页面设置API参数")
        st.stop()

    # 模板名称
    template_name = st.text_input("模板名称", placeholder="例如：学术论文模板")
    template_desc = st.text_input("模板描述", placeholder="例如：适用于学术论文排版")

    # 格式要求文本
    format_text = st.text_area(
        "格式要求描述",
        height=200,
        placeholder="请输入格式要求，例如：\n"
                    "标题使用黑体小二号字，居中对齐\n"
                    "正文使用宋体小四号字，左对齐，行距1.5倍\n"
                    "一级标题使用黑体三号字，左对齐"
    )

    if st.button("🪄 AI解析生成模板", type="primary"):
        if not template_name:
            st.error("请输入模板名称")
        elif not format_text:
            st.error("请输入格式要求描述")
        else:
            with st.spinner("AI正在解析..."):
                try:
                    # 构建解析提示词
                    prompt = f"""请根据以下格式要求描述，生成一个JSON格式的排版模板。

格式要求：
{format_text}

请严格按照以下JSON格式输出：
{{
    "name": "{template_name}",
    "description": "{template_desc or template_name}",
    "rules": {{
        "标题": {{"font": "黑体", "size": "小二", "alignment": "center", "bold": true}},
        "一级标题": {{"font": "黑体", "size": "三号", "alignment": "left", "bold": true}},
        "二级标题": {{"font": "黑体", "size": "小三", "alignment": "left", "bold": true}},
        "正文": {{"font": "宋体", "size": "小四", "alignment": "justify", "bold": false}}
    }}
}}

只输出JSON，不要输出其他内容。"""

                    # 调用AI
                    api_config = {
                        "api_url": st.session_state.api_url,
                        "api_key": st.session_state.api_key,
                        "model": st.session_state.model,
                        "timeout": 120
                    }
                    connector = AIConnector(api_config)

                    success, response = connector.send_request(prompt)
                    if not success:
                        st.error(f"AI请求失败: {response}")
                        return

                    # 解析响应
                    success, result = connector.parse_response(response)
                    if not success:
                        st.error(f"解析失败: {result}")
                        return

                    # 保存模板
                    format_manager = FormatManager()
                    if format_manager.save_template(template_name, result):
                        st.success(f"✅ 模板已创建: {template_name}")
                        st.json(result)
                    else:
                        st.error("保存模板失败")

                except Exception as e:
                    st.error(f"处理失败: {str(e)}")


# ========================
# 页面：帮助
# ========================
def page_help():
    """帮助页面"""
    st.markdown('<p class="main-header">❓ 帮助说明</p>', unsafe_allow_html=True)

    st.markdown("""
    ## FormulaAI Web - AI智能文档排版工具

    ### 功能介绍

    FormulaAI 是一款基于AI的智能文档排版工具，能够自动识别文档结构并应用专业排版格式。

    ### 使用步骤

    1. **配置API**
       - 在「API配置」页面设置您的AI API参数
       - 支持DeepSeek、OpenAI等兼容API
       - 点击「测试连接」验证配置

    2. **管理模板**
       - 在「模板管理」页面创建或编辑排版模板
       - 定义各类文档元素（标题、正文等）的格式规则

    3. **排版文档**
       - 在「文档排版」页面上传Word文档
       - 选择排版模板
       - 点击「开始排版」
       - 下载排版后的文档

    4. **智能解析**（可选）
       - 在「从文本解析」页面输入自然语言描述
       - AI自动生成排版模板

    ### 支持的API

    - **DeepSeek**: `https://api.deepseek.com`
    - **OpenAI**: `https://api.openai.com`
    - **其他兼容API**: 输入对应的Base URL

    ### 常见问题

    **Q: 排版后格式不对？**
    A: 请检查模板规则是否正确设置，或尝试使用不同的AI模型。

    **Q: API调用失败？**
    A: 请确认API URL和Key正确，检查网络连接，确保账户余额充足。

    **Q: 支持哪些文档格式？**
    A: 目前仅支持 .docx 格式的Word文档。

    ### 版本信息

    - 版本: 2.0.0 (Web版)
    - 框架: Streamlit
    - 核心模块: AIConnector, FormatManager, DocProcessor
    """)

    st.divider()
    st.markdown("""
    <div style="text-align: center; color: #666;">
        FormulaAI Web v2.0.0 |
        <a href="https://github.com/waterdrop26651/FormulaAI" target="_blank">GitHub</a>
    </div>
    """, unsafe_allow_html=True)


# ========================
# 主应用
# ========================
def main():
    """主应用"""
    # 侧边栏导航
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/document.png", width=80)
        st.title("FormulaAI")
        st.caption("AI智能文档排版工具")

        st.divider()

        page = st.radio(
            "导航",
            ["📄 文档排版", "📝 模板管理", "⚙️ API配置", "🔮 从文本解析", "❓ 帮助"],
            label_visibility="collapsed"
        )

        st.divider()

        # 状态显示
        if st.session_state.api_url and st.session_state.api_key:
            st.success("✅ API已配置")
        else:
            st.warning("⚠️ API未配置")

    # 根据选择显示页面
    if page == "📄 文档排版":
        page_document_format()
    elif page == "📝 模板管理":
        page_template_management()
    elif page == "⚙️ API配置":
        page_api_config()
    elif page == "🔮 从文本解析":
        page_text_parsing()
    elif page == "❓ 帮助":
        page_help()


if __name__ == "__main__":
    main()
