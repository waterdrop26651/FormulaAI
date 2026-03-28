# -*- coding: utf-8 -*-
"""Tests for the runtime document-format harness."""

import json
from pathlib import Path

from docx import Document

from src.runtime.contracts import RunStage, RunStatus, RuntimeErrorCode
from src.runtime.document_format_harness import DocumentFormatHarness


class FakeFormatManager:
    def __init__(self, template=None):
        self.template = template or {
            "name": "测试模板",
            "rules": {
                "标题": {"font": "黑体", "size": "小二", "bold": True, "alignment": "居中", "line_spacing": 1.0},
                "正文": {"font": "宋体", "size": "小四", "bold": False, "alignment": "两端对齐", "line_spacing": 1.5},
            },
        }

    def get_template(self, template_name):
        return self.template if template_name == "测试模板" else None


class FakeAIConnector:
    def __init__(self, api_config):
        self.api_config = api_config

    def validate_config(self):
        return True, "ok"

    def generate_prompt(self, paragraphs, rules):
        return json.dumps({"paragraphs": paragraphs, "rules": rules}, ensure_ascii=False)

    def send_request(self, prompt):
        return True, {
            "choices": [
                {
                    "message": {
                        "content": json.dumps(
                            {
                                "elements": [
                                    {
                                        "type": "标题",
                                        "content": "论文标题",
                                        "format": {"font": "黑体", "size": "小二", "bold": True, "line_spacing": 1.0, "alignment": "center"},
                                    },
                                    {
                                        "type": "正文",
                                        "content": "这是正文。",
                                        "format": {"font": "宋体", "size": "小四", "bold": False, "line_spacing": 1.5, "alignment": "justify"},
                                    },
                                ]
                            },
                            ensure_ascii=False,
                        )
                    }
                }
            ]
        }

    def parse_response(self, response):
        return True, json.loads(response["choices"][0]["message"]["content"])


class InvalidConfigAIConnector(FakeAIConnector):
    def validate_config(self):
        return False, "missing url"


class ReportingDocProcessor:
    def __init__(self):
        self.input_file = None
        self.output_file = None

    def read_document(self, file_path):
        self.input_file = file_path
        return True

    def get_document_text(self):
        return ["论文标题", "这是正文。"]

    def apply_formatting(self, formatting_instructions, custom_save_path=None, header_footer_config=None):
        output_path = Path(custom_save_path) / "output.docx"
        document = Document()
        for element in formatting_instructions["elements"]:
            document.add_paragraph(element["content"])
        document.save(output_path)
        self.output_file = str(output_path)
        return {
            "success": True,
            "total_elements": len(formatting_instructions["elements"]),
            "processed_elements": len(formatting_instructions["elements"]),
            "failed_elements": [],
            "warnings": [],
            "backup_path": None,
            "output_file": self.output_file,
            "header_footer": {"attempted": True, "success": True, "error": None},
        }

    def get_output_file(self):
        return self.output_file


class FailedReportingDocProcessor(ReportingDocProcessor):
    def apply_formatting(self, formatting_instructions, custom_save_path=None, header_footer_config=None):
        return {
            "success": False,
            "total_elements": len(formatting_instructions["elements"]),
            "processed_elements": 1,
            "failed_elements": [{"index": 1, "error": "boom"}],
            "warnings": ["partial failure"],
            "backup_path": None,
            "output_file": None,
            "header_footer": {"attempted": True, "success": True, "error": None},
        }


def _docx_bytes():
    document = Document()
    document.add_paragraph("论文标题")
    document.add_paragraph("这是正文。")
    path = "/tmp/runtime_harness_input.docx"
    document.save(path)
    with open(path, "rb") as handle:
        return handle.read()


def test_document_format_harness_happy_path(tmp_path):
    harness = DocumentFormatHarness(
        runtime_dir=tmp_path / "runtime",
        format_manager=FakeFormatManager(),
        doc_processor_factory=ReportingDocProcessor,
        ai_connector_factory=FakeAIConnector,
    )

    result = harness.run(
        source_name="input.docx",
        source_bytes=_docx_bytes(),
        template_name="测试模板",
        api_config={"api_url": "https://example.com", "api_key": "k", "model": "demo", "timeout": 1},
        header_footer_config={},
    )

    assert result.status == RunStatus.SUCCEEDED
    assert result.final_stage == RunStage.COMPLETED
    assert result.output_bytes
    assert result.instruction_count == 2
    assert result.output_path is None
    assert result.stage_history[0].stage == RunStage.INIT
    assert result.render_report["processed_elements"] == 2
    manifests = list((tmp_path / "runtime" / "runs").rglob("manifest.json"))
    assert manifests
    manifest = json.loads(manifests[0].read_text(encoding="utf-8"))
    assert manifest["status"] == "success"
    assert manifest["result"]["element_count"] == 2
    assert manifest["result"]["output_sha256"].startswith("sha256:")
    assert manifest["result"]["render_summary"]["processed_elements"] == 2
    assert manifest["output_validation"]["document_loadable"] is True
    assert manifest["output_validation"]["paragraph_count"] == 2


def test_document_format_harness_invalid_api_config_returns_failed_result(tmp_path):
    harness = DocumentFormatHarness(
        runtime_dir=tmp_path / "runtime",
        format_manager=FakeFormatManager(),
        doc_processor_factory=ReportingDocProcessor,
        ai_connector_factory=InvalidConfigAIConnector,
    )

    result = harness.run(
        source_name="input.docx",
        source_bytes=_docx_bytes(),
        template_name="测试模板",
        api_config={"api_url": "", "api_key": "", "model": "", "timeout": 1},
        header_footer_config={},
    )

    assert result.status == RunStatus.FAILED
    assert result.final_stage == RunStage.FAILED
    assert result.error_code == RuntimeErrorCode.INVALID_API_CONFIG
    assert result.error_message == "missing url"
    manifests = list((tmp_path / "runtime" / "runs").rglob("manifest.json"))
    failure_files = list((tmp_path / "runtime" / "runs").rglob("failure.json"))
    assert manifests and failure_files
    manifest = json.loads(manifests[0].read_text(encoding="utf-8"))
    failure = json.loads(failure_files[0].read_text(encoding="utf-8"))
    assert manifest["error"]["code"] == RuntimeErrorCode.INVALID_API_CONFIG.value
    assert failure["message"] == "missing url"
    assert "INVALID_API_CONFIG" not in failure["message"]


def test_document_format_harness_missing_template_returns_failed_result(tmp_path):
    harness = DocumentFormatHarness(
        runtime_dir=tmp_path / "runtime",
        format_manager=FakeFormatManager(template=None),
        doc_processor_factory=ReportingDocProcessor,
        ai_connector_factory=FakeAIConnector,
    )

    result = harness.run(
        source_name="input.docx",
        source_bytes=_docx_bytes(),
        template_name="不存在模板",
        api_config={"api_url": "https://example.com", "api_key": "k", "model": "demo", "timeout": 1},
        header_footer_config={},
    )

    assert result.status == RunStatus.FAILED
    assert result.error_code == RuntimeErrorCode.TEMPLATE_NOT_FOUND


def test_document_format_harness_invalid_header_footer_returns_failed_result(tmp_path):
    harness = DocumentFormatHarness(
        runtime_dir=tmp_path / "runtime",
        format_manager=FakeFormatManager(),
        doc_processor_factory=ReportingDocProcessor,
        ai_connector_factory=FakeAIConnector,
    )

    result = harness.run(
        source_name="input.docx",
        source_bytes=_docx_bytes(),
        template_name="测试模板",
        api_config={"api_url": "https://example.com", "api_key": "k", "model": "demo", "timeout": 1},
        header_footer_config={"header_alignment": "bad-value"},
    )

    assert result.status == RunStatus.FAILED
    assert result.error_code == RuntimeErrorCode.HEADER_FOOTER_INVALID


def test_document_format_harness_surfaces_failed_render_report(tmp_path):
    harness = DocumentFormatHarness(
        runtime_dir=tmp_path / "runtime",
        format_manager=FakeFormatManager(),
        doc_processor_factory=FailedReportingDocProcessor,
        ai_connector_factory=FakeAIConnector,
    )

    result = harness.run(
        source_name="input.docx",
        source_bytes=_docx_bytes(),
        template_name="测试模板",
        api_config={"api_url": "https://example.com", "api_key": "k", "model": "demo", "timeout": 1},
        header_footer_config={},
    )

    assert result.status == RunStatus.FAILED
    assert result.error_code == RuntimeErrorCode.FORMATTING_FAILED
    assert result.render_report["failed_elements"]
