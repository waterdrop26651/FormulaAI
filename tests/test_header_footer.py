# -*- coding: utf-8 -*-
"""Tests for header/footer config and processor."""

from docx import Document

from src.core.header_footer_config import HeaderFooterConfig
from src.core.header_footer_processor import HeaderFooterProcessor


def test_header_footer_config_roundtrip_and_validation():
    config = HeaderFooterConfig(
        enable_header=True,
        enable_footer=True,
        header_content="测试页眉",
        footer_content="测试页脚",
        include_page_number=True,
        page_number_position="footer_center",
        page_number_format="第 {page} 页",
    )

    valid, message = config.validate()
    assert valid is True
    assert "通过" in message

    restored = HeaderFooterConfig.from_dict(config.to_dict())
    assert restored.header_content == "测试页眉"
    assert restored.footer_content == "测试页脚"

    restored.page_number_format = "第 页"
    valid, _ = restored.validate()
    assert valid is False


def test_header_footer_config_three_column_effective_content():
    config = HeaderFooterConfig(
        use_three_column_layout=True,
        header_left_content="L",
        header_center_content="C",
        header_right_content="R",
        footer_left_content="A",
        footer_center_content="B",
        footer_right_content="D",
    )

    assert config.get_effective_header_content() == "L\tC\tR"
    assert config.get_effective_footer_content() == "A\tB\tD"


def test_header_footer_processor_applies_header_footer_and_page_number():
    document = Document()
    document.add_paragraph("正文测试")
    processor = HeaderFooterProcessor()
    config = HeaderFooterConfig(
        enable_header=True,
        enable_footer=True,
        header_content="页眉内容",
        footer_content="页脚内容",
        include_page_number=True,
        page_number_position="footer_right",
        page_number_format="第 {page} 页",
    )

    success = processor.apply_header_footer(document, config)
    section = document.sections[0]

    assert success is True
    assert "页眉内容" in section.header.paragraphs[0].text
    assert "页脚内容" in section.footer.paragraphs[0].text
    assert "第 1 页" in section.footer.paragraphs[0].text


def test_header_footer_processor_rejects_invalid_config():
    document = Document()
    processor = HeaderFooterProcessor()
    config = HeaderFooterConfig(header_alignment="invalid")

    assert processor.apply_header_footer(document, config) is False
