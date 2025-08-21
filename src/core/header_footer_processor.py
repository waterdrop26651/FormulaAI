# -*- coding: utf-8 -*-
"""
页眉页脚处理模块
负责将页眉页脚配置应用到Word文档中
"""

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ..utils.logger import app_logger
from .header_footer_config import HeaderFooterConfig

class HeaderFooterProcessor:
    """页眉页脚处理器，负责应用页眉页脚设置到文档"""
    
    def __init__(self):
        """初始化页眉页脚处理器"""
        self.logger = app_logger
        
        # 对齐方式映射
        self.alignment_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        
        app_logger.debug("页眉页脚处理器初始化完成")
    
    def apply_header_footer(self, document, config: HeaderFooterConfig) -> bool:
        """
        应用页眉页脚设置到文档
        
        Args:
            document: python-docx Document对象
            config: 页眉页脚配置
            
        Returns:
            bool: 是否成功应用
        """
        try:
            # 验证配置
            is_valid, error_msg = config.validate()
            if not is_valid:
                self.logger.error(f"页眉页脚配置无效: {error_msg}")
                return False
            
            # 获取第一个section（大多数情况下只有一个section）
            if not document.sections:
                self.logger.error("文档没有section，无法设置页眉页脚")
                return False
            
            section = document.sections[0]
            
            # 设置高级选项
            self._configure_advanced_settings(document, section, config)
            
            # 应用页眉
            if config.enable_header:
                success = self._apply_header(section, config)
                if not success:
                    self.logger.warning("页眉应用失败，但继续处理页脚")
            
            # 应用页脚
            if config.enable_footer:
                success = self._apply_footer(section, config)
                if not success:
                    self.logger.warning("页脚应用失败")
            
            # 应用页码
            if config.include_page_number:
                self._apply_page_numbers(section, config)
            
            self.logger.info("页眉页脚应用成功")
            return True
            
        except Exception as e:
            self.logger.error(f"应用页眉页脚失败: {str(e)}")
            import traceback
            self.logger.error(f"详细错误: {traceback.format_exc()}")
            return False
    
    def _configure_advanced_settings(self, document, section, config: HeaderFooterConfig):
        """配置高级设置"""
        try:
            # 首页不同
            if config.different_first_page:
                section.different_first_page_header_footer = True
                self.logger.debug("启用首页不同页眉页脚")
            
            # 奇偶页不同
            if config.different_odd_even:
                document.settings.odd_and_even_pages_header_footer = True
                self.logger.debug("启用奇偶页不同页眉页脚")
                
        except Exception as e:
            self.logger.error(f"配置高级设置失败: {str(e)}")
    
    def _apply_header(self, section, config: HeaderFooterConfig) -> bool:
        """应用页眉设置"""
        try:
            # 普通页眉
            header = section.header
            header.is_linked_to_previous = False
            
            content = config.get_effective_header_content()
            if content.strip():  # 只有在有内容时才设置
                self._set_header_footer_content(header, content, config, is_header=True)
                self.logger.debug(f"设置普通页眉: {content[:50]}...")
            
            # 首页页眉
            if config.different_first_page and config.first_page_header_content.strip():
                first_header = section.first_page_header
                first_header.is_linked_to_previous = False
                self._set_header_footer_content(first_header, config.first_page_header_content, config, is_header=True)
                self.logger.debug(f"设置首页页眉: {config.first_page_header_content[:50]}...")
            
            # 偶数页页眉
            if config.different_odd_even and config.even_page_header_content.strip():
                even_header = section.even_page_header
                even_header.is_linked_to_previous = False
                self._set_header_footer_content(even_header, config.even_page_header_content, config, is_header=True)
                self.logger.debug(f"设置偶数页页眉: {config.even_page_header_content[:50]}...")
            
            return True
            
        except Exception as e:
            self.logger.error(f"应用页眉失败: {str(e)}")
            return False
    
    def _apply_footer(self, section, config: HeaderFooterConfig) -> bool:
        """应用页脚设置"""
        try:
            # 普通页脚
            footer = section.footer
            footer.is_linked_to_previous = False
            
            content = config.get_effective_footer_content()
            if content.strip():  # 只有在有内容时才设置
                self._set_header_footer_content(footer, content, config, is_header=False)
                self.logger.debug(f"设置普通页脚: {content[:50]}...")
            
            # 首页页脚
            if config.different_first_page and config.first_page_footer_content.strip():
                first_footer = section.first_page_footer
                first_footer.is_linked_to_previous = False
                self._set_header_footer_content(first_footer, config.first_page_footer_content, config, is_header=False)
                self.logger.debug(f"设置首页页脚: {config.first_page_footer_content[:50]}...")
            
            # 偶数页页脚
            if config.different_odd_even and config.even_page_footer_content.strip():
                even_footer = section.even_page_footer
                even_footer.is_linked_to_previous = False
                self._set_header_footer_content(even_footer, config.even_page_footer_content, config, is_header=False)
                self.logger.debug(f"设置偶数页页脚: {config.even_page_footer_content[:50]}...")
            
            return True
            
        except Exception as e:
            self.logger.error(f"应用页脚失败: {str(e)}")
            return False
    
    def _set_header_footer_content(self, header_footer_obj, content: str, config: HeaderFooterConfig, is_header: bool = True):
        """设置页眉或页脚的内容和格式"""
        try:
            # 获取第一个段落
            para = header_footer_obj.paragraphs[0]
            para.text = content
            
            # 设置对齐方式
            alignment = config.header_alignment if is_header else config.footer_alignment
            para.alignment = self.alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
            
            # 设置字体格式
            if para.runs:
                run = para.runs[0]
            else:
                run = para.add_run()
            
            # 应用字体设置
            font_name = config.header_font if is_header else config.footer_font
            font_size = config.header_font_size if is_header else config.footer_font_size
            font_bold = config.header_bold if is_header else config.footer_bold
            font_italic = config.header_italic if is_header else config.footer_italic
            
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = font_bold
            run.font.italic = font_italic
            
            # 尝试应用样式（如果存在）
            try:
                style_name = "Header" if is_header else "Footer"
                if hasattr(header_footer_obj, 'part') and hasattr(header_footer_obj.part, 'document'):
                    para.style = header_footer_obj.part.document.styles[style_name]
                    self.logger.debug(f"应用{style_name}样式成功")
            except (KeyError, AttributeError):
                # 样式不存在时忽略
                self.logger.debug(f"{'页眉' if is_header else '页脚'}样式不存在，使用默认格式")
                
        except Exception as e:
            self.logger.error(f"设置{'页眉' if is_header else '页脚'}内容失败: {str(e)}")
    
    def _apply_page_numbers(self, section, config: HeaderFooterConfig):
        """应用页码设置"""
        try:
            position = config.page_number_position
            page_text = config.page_number_format.replace("{page}", "1")  # 占位符，实际页码由Word自动处理
            
            if "header" in position:
                header = section.header
                header.is_linked_to_previous = False
                self._add_page_number_to_header_footer(header, page_text, position)
                self.logger.debug(f"在页眉添加页码: {position}")
            elif "footer" in position:
                footer = section.footer
                footer.is_linked_to_previous = False
                self._add_page_number_to_header_footer(footer, page_text, position)
                self.logger.debug(f"在页脚添加页码: {position}")
                
        except Exception as e:
            self.logger.error(f"应用页码失败: {str(e)}")
    
    def _add_page_number_to_header_footer(self, header_footer_obj, page_text: str, position: str):
        """向页眉或页脚添加页码"""
        try:
            para = header_footer_obj.paragraphs[0]
            
            # 根据位置设置对齐方式
            if "left" in position:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif "center" in position:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif "right" in position:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # 如果已有内容，添加到现有内容
            if para.text.strip():
                # 根据位置决定页码的位置
                if "left" in position:
                    para.text = f"{page_text}\t{para.text}"
                elif "right" in position:
                    para.text = f"{para.text}\t{page_text}"
                else:  # center
                    para.text = f"{para.text}\t{page_text}"
            else:
                para.text = page_text
                
        except Exception as e:
            self.logger.error(f"添加页码失败: {str(e)}")
    
    def remove_header_footer(self, document) -> bool:
        """移除文档的页眉页脚"""
        try:
            for section in document.sections:
                # 移除页眉
                header = section.header
                header.is_linked_to_previous = True  # 链接到前一节，实际上是删除
                
                # 移除页脚
                footer = section.footer
                footer.is_linked_to_previous = True  # 链接到前一节，实际上是删除
                
                # 移除首页页眉页脚
                if section.different_first_page_header_footer:
                    first_header = section.first_page_header
                    first_header.is_linked_to_previous = True
                    
                    first_footer = section.first_page_footer
                    first_footer.is_linked_to_previous = True
                
                # 移除偶数页页眉页脚
                try:
                    if document.settings.odd_and_even_pages_header_footer:
                        even_header = section.even_page_header
                        even_header.is_linked_to_previous = True
                        
                        even_footer = section.even_page_footer
                        even_footer.is_linked_to_previous = True
                except AttributeError:
                    # 某些版本可能不支持
                    pass
            
            self.logger.info("页眉页脚移除成功")
            return True
            
        except Exception as e:
            self.logger.error(f"移除页眉页脚失败: {str(e)}")
            return False