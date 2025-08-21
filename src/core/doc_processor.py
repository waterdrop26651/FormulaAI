# -*- coding: utf-8 -*-
"""
文档处理模块
负责读取、解析和写入Word文档，以及应用排版格式。
"""

import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from ..utils.logger import app_logger
from ..utils.file_utils import generate_output_filename, is_valid_docx, backup_file
from ..utils.font_manager import FontManager

class DocProcessor:
    """文档处理器，负责读取、解析和写入Word文档"""
    
    def __init__(self):
        """初始化文档处理器"""
        self.document = None
        self.input_file = None
        self.output_file = None
        self.paragraphs_text = []
        
        # 使用字体管理器获取字体信息
        self.font_manager = FontManager()
        
        # 字号映射
        self.font_size_mapping = {
            "小二": Pt(18),
            "三号": Pt(16),
            "小三": Pt(15),
            "四号": Pt(14),
            "小四": Pt(12),
            "五号": Pt(10.5),
            "小五": Pt(9),
            "六号": Pt(7.5)
        }
        
        app_logger.debug(f"文档处理器初始化完成，字号映射: {list(self.font_size_mapping.keys())}")
    
    def read_document(self, file_path):
        """
        读取Word文档内容
        
        Args:
            file_path: 文档路径
            
        Returns:
            是否成功读取文档
        """
        if not is_valid_docx(file_path):
            app_logger.error(f"无效的Word文档: {file_path}")
            return False
        
        try:
            self.document = Document(file_path)
            self.input_file = file_path
            self.paragraphs_text = [p.text for p in self.document.paragraphs]
            
            app_logger.info(f"成功读取文档: {file_path}")
            app_logger.debug(f"文档包含 {len(self.document.paragraphs)} 个段落")
            return True
        except Exception as e:
            app_logger.error(f"读取文档失败: {file_path}, 错误: {str(e)}")
            return False
    
    def get_document_text(self):
        """
        获取文档的纯文本内容
        
        Returns:
            文档的段落文本列表
        """
        if not self.document:
            app_logger.warning("尚未加载文档")
            return []
        
        return self.paragraphs_text
    
    def apply_formatting(self, formatting_instructions, custom_save_path=None):
        """
        根据排版指令应用格式
        
        Args:
            formatting_instructions: 排版指令，包含元素类型和格式信息
            custom_save_path: 自定义保存路径，如果指定则使用该路径
            
        Returns:
            是否成功应用格式
        """
        if not self.document:
            app_logger.error("尚未加载文档，无法应用格式")
            return False
        
        try:
            # 记录格式化指令，方便调试
            app_logger.debug(f"格式化指令: {formatting_instructions}")
            
            # 检查formatting_instructions结构
            if not isinstance(formatting_instructions, dict):
                app_logger.error(f"排版指令不是字典类型: {type(formatting_instructions)}")
                return False
            
            elements = formatting_instructions.get('elements', [])
            if not isinstance(elements, list):
                app_logger.error(f"元素列表不是列表类型: {type(elements)}")
                return False
            
            app_logger.info(f"开始应用排版格式，共有 {len(elements)} 个元素")
            
            # 记录自定义保存路径
            if custom_save_path:
                app_logger.info(f"使用自定义保存路径: {custom_save_path}")
            
            # 备份原始文档
            backup_path = backup_file(self.input_file)
            app_logger.info(f"文件备份成功: {self.input_file} -> {backup_path}")
            
            # 创建新文档以应用格式
            try:
                new_doc = Document()
                app_logger.debug("成功创建新文档")
            except Exception as e:
                app_logger.error(f"创建新文档失败: {str(e)}")
                return False
            
            # 批处理元素，避免内存压力
            batch_size = 10  # 每批处理10个元素，进一步减少内存压力
            total_elements = len(elements)
            
            for batch_start in range(0, total_elements, batch_size):
                batch_end = min(batch_start + batch_size, total_elements)
                app_logger.info(f"处理批次: {batch_start+1}-{batch_end}/{total_elements}")
                
                # 处理当前批次的元素
                for i in range(batch_start, batch_end):
                    element = elements[i]
                    try:
                        app_logger.debug(f"处理第 {i+1}/{total_elements} 个元素")
                        self._process_element(new_doc, element)
                    except Exception as e:
                        app_logger.error(f"处理第 {i+1} 个元素时发生错误: {str(e)}")
                        # 继续处理下一个元素，不中断整个过程
                
                # 批次处理完成后，强制垃圾回收
                import gc
                gc.collect()
                app_logger.debug(f"批次 {batch_start+1}-{batch_end} 处理完成，已执行垃圾回收")
            
            # 生成输出文件名
            try:
                self.output_file = generate_output_filename(self.input_file, custom_save_path)
                app_logger.debug(f"生成输出文件名: {self.output_file}")
            except Exception as e:
                app_logger.error(f"生成输出文件名失败: {str(e)}")
                # 如果指定了自定义保存路径，则使用该路径
                if custom_save_path and os.path.isdir(custom_save_path):
                    filename = os.path.basename(self.input_file).replace('.docx', '_已排版.docx')
                    self.output_file = os.path.join(custom_save_path, filename)
                else:
                    self.output_file = self.input_file.replace('.docx', '_已排版.docx')
                app_logger.debug(f"使用默认输出文件名: {self.output_file}")
            
            # 保存文档
            try:
                # 确保输出目录存在
                output_dir = os.path.dirname(self.output_file)
                if not os.path.exists(output_dir):
                    os.makedirs(output_dir)
                    app_logger.debug(f"创建输出目录: {output_dir}")
                
                new_doc.save(self.output_file)
                app_logger.info(f"成功应用排版格式并保存到: {self.output_file}")
                return True
            except Exception as e:
                app_logger.error(f"保存文档失败: {str(e)}")
                return False
        except Exception as e:
            app_logger.error(f"应用排版格式失败: {str(e)}")
            import traceback
            app_logger.error(f"详细错误信息: {traceback.format_exc()}")
            return False
    
    def _process_element(self, doc, element):
        """
        处理单个排版元素
        
        Args:
            doc: 目标文档对象
            element: 排版元素信息
        """
        paragraph = None
        run = None
        
        try:
            # 记录当前处理的元素信息
            app_logger.debug(f"开始处理元素: {element}")
            
            # 检查元素结构
            if not isinstance(element, dict):
                app_logger.error(f"元素不是字典类型: {type(element)}")
                return
            
            content = element.get('content', '')
            app_logger.debug(f"元素内容: {content[:50]}...")
            
            element_type = element.get('type', '正文')
            app_logger.debug(f"元素类型: {element_type}")
            
            format_info = element.get('format', {})
            app_logger.debug(f"格式信息: {format_info}")
            
            # 检查format_info结构
            if not isinstance(format_info, dict):
                app_logger.error(f"格式信息不是字典类型: {type(format_info)}")
                format_info = {}
            
            # 添加段落
            paragraph = doc.add_paragraph()
            app_logger.debug("成功添加段落")
            
            run = paragraph.add_run(content)
            app_logger.debug("成功添加文本运行")
            
            # 应用字体
            self._apply_font(run, format_info)
            app_logger.debug("成功应用字体格式")
            
            # 应用段落格式
            self._apply_paragraph_format(paragraph, format_info, element_type)
            app_logger.debug("成功应用段落格式")
            
            app_logger.debug(f"完成元素处理: {element_type}, 内容: {content[:20]}...")
            
        except Exception as e:
            app_logger.error(f"处理元素时发生异常: {str(e)}")
            # 记录详细错误信息
            import traceback
            app_logger.error(f"异常详情: {traceback.format_exc()}")
            # 不抛出异常，继续处理下一个元素
        finally:
            # 清理局部变量，释放内存
            try:
                del paragraph, run
            except:
                pass
            
            # 强制垃圾回收，防止内存泄漏
            import gc
            gc.collect()
    
    def _apply_font(self, run, format_info):
        """
        应用字体格式，增强内存安全性
        
        Args:
            run: 文本运行对象
            format_info: 格式信息
        """
        if not run or not hasattr(run, 'font'):
            app_logger.error("无效的文本运行对象")
            return
            
        try:
            # 字体名称 - 使用安全的默认值
            font_name = format_info.get('font', '宋体') if format_info else '宋体'
            app_logger.debug(f"原始字体名称: {font_name}")
            
            # 简化字体处理，避免复杂的映射操作
            safe_fonts = {'宋体': 'SimSun', '黑体': 'SimHei', '楷体': 'KaiTi', '仿宋': 'FangSong'}
            document_font = safe_fonts.get(font_name, font_name)
            app_logger.debug(f"用于文档的字体名称: {document_font}")
            
            # 安全设置字体名称
            try:
                if hasattr(run.font, 'name'):
                    run.font.name = document_font
            except Exception as e:
                app_logger.error(f"设置字体名称失败: {str(e)}")
            
            # 简化中文字体设置，避免直接操作XML元素
            try:
                if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), document_font)
                    app_logger.debug(f"成功设置中文字体: {document_font}")
            except Exception as e:
                app_logger.debug(f"设置中文字体失败，使用默认设置: {str(e)}")
                
            # 字体大小 - 使用安全的默认值
            font_size = format_info.get('size', '小四') if format_info else '小四'
            app_logger.debug(f"原始字体大小: {font_size}")
            
            try:
                if isinstance(font_size, str) and font_size in self.font_size_mapping:
                    mapped_size = self.font_size_mapping[font_size]
                    app_logger.debug(f"映射后的字体大小: {mapped_size}")
                    if hasattr(run.font, 'size'):
                        run.font.size = mapped_size
            except Exception as e:
                app_logger.error(f"设置字体大小失败: {str(e)}")
            
            # 安全设置字体属性
            try:
                bold = format_info.get('bold', False) if format_info else False
                if hasattr(run, 'bold'):
                    run.bold = bold
                    app_logger.debug(f"设置粗体: {bold}")
            except Exception as e:
                app_logger.error(f"设置粗体失败: {str(e)}")
            
            try:
                italic = format_info.get('italic', False) if format_info else False
                if hasattr(run, 'italic'):
                    run.italic = italic
                    app_logger.debug(f"设置斜体: {italic}")
            except Exception as e:
                app_logger.error(f"设置斜体失败: {str(e)}")
            
            try:
                underline = format_info.get('underline', False) if format_info else False
                if hasattr(run, 'underline'):
                    run.underline = underline
                    app_logger.debug(f"设置下划线: {underline}")
            except Exception as e:
                app_logger.error(f"设置下划线失败: {str(e)}")
            
        except Exception as e:
            app_logger.error(f"应用字体格式时发生严重异常: {str(e)}")
            import traceback
            app_logger.error(f"异常详情: {traceback.format_exc()}")
            # 不抛出异常，使用默认字体设置
    
    def _apply_paragraph_format(self, paragraph, format_info, element_type):
        """
        应用段落格式，增强内存安全性
        
        Args:
            paragraph: 段落对象
            format_info: 格式信息
            element_type: 元素类型
        """
        if not paragraph or not hasattr(paragraph, 'paragraph_format'):
            app_logger.error("无效的段落对象")
            return
            
        try:
            # 安全获取格式信息
            if not format_info or not isinstance(format_info, dict):
                format_info = {}
            
            # 行间距 - 使用更安全的默认值
            line_spacing = format_info.get('line_spacing', 1.5)
            app_logger.debug(f"设置行间距: {line_spacing}")
            
            try:
                if isinstance(line_spacing, (int, float)) and hasattr(paragraph, 'paragraph_format'):
                    # 防止异常大的行间距值导致程序崩溃
                    if line_spacing > 3.0 or line_spacing < 0.8:
                        app_logger.warning(f"检测到异常行间距值: {line_spacing}，将使用默认值1.5")
                        line_spacing = 1.5
                    
                    if abs(line_spacing - 1.0) < 0.1:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        app_logger.debug("设置单倍行间距")
                    elif abs(line_spacing - 1.5) < 0.1:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        app_logger.debug("设置1.5倍行间距")
                    elif abs(line_spacing - 2.0) < 0.1:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                        app_logger.debug("设置双倍行间距")
                    else:
                        # 使用默认1.5倍行间距，避免复杂设置
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        app_logger.debug("使用默认1.5倍行间距")
            except Exception as e:
                app_logger.error(f"设置行间距失败: {str(e)}")
            
            # 对齐方式 - 简化处理
            try:
                alignment = format_info.get('alignment', 'left')
                app_logger.debug(f"设置对齐方式: {alignment}")
                
                if hasattr(paragraph, 'alignment'):
                    if alignment == 'center':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif alignment == 'right':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif alignment == 'justify':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    else:  # 默认左对齐
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            except Exception as e:
                app_logger.error(f"设置对齐方式失败: {str(e)}")
            
            # 首行缩进 - 简化处理
            try:
                if hasattr(paragraph, 'paragraph_format') and hasattr(paragraph.paragraph_format, 'first_line_indent'):
                    first_line_indent = format_info.get('first_line_indent', None)
                    
                    if first_line_indent is not None and isinstance(first_line_indent, (int, float)):
                        if 0 <= first_line_indent <= 50:  # 限制缩进范围
                            paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)
                            app_logger.debug(f"设置首行缩进: {first_line_indent}磅")
                    elif element_type == '正文':  # 正文默认缩进
                        paragraph.paragraph_format.first_line_indent = Pt(21)
                        app_logger.debug("设置正文默认首行缩进: 21磅")
            except Exception as e:
                app_logger.error(f"设置首行缩进失败: {str(e)}")
            
            # 段间距 - 简化处理，避免复杂设置
            try:
                if hasattr(paragraph, 'paragraph_format'):
                    # 使用固定的安全间距值
                    if hasattr(paragraph.paragraph_format, 'space_before'):
                        paragraph.paragraph_format.space_before = Pt(0)
                    if hasattr(paragraph.paragraph_format, 'space_after'):
                        paragraph.paragraph_format.space_after = Pt(0)
                    app_logger.debug("设置段间距为0")
            except Exception as e:
                app_logger.error(f"设置段间距失败: {str(e)}")
                
        except Exception as e:
            app_logger.error(f"应用段落格式时发生严重异常: {str(e)}")
            import traceback
            app_logger.error(f"异常详情: {traceback.format_exc()}")
            # 不抛出异常，使用默认段落设置
    
    def get_output_file(self):
        """
        获取输出文件路径
        
        Returns:
            输出文件路径
        """
        return self.output_file
