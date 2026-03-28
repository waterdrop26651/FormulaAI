# -*- coding: utf-8 -*-
"""Offline eval harness for replayable runtime verification."""

import argparse
import io
import json
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from docx import Document

from src.core.ai_connector import AIConnector
from src.core.doc_processor import DocProcessor
from src.runtime.document_format_harness import DocumentFormatHarness


def _iso_now():
    return datetime.now(timezone.utc).isoformat()


def _enum_value(value):
    return value.value if hasattr(value, "value") else value


class ReplayAIConnector(AIConnector):
    """Offline connector that replays fixed validation and response payloads."""

    def __init__(self, api_config, ai_case):
        merged_config = {
            "api_url": "https://eval.local/v1/chat/completions",
            "api_key": "eval-key",
            "model": "eval-model",
            "timeout": 1,
        }
        merged_config.update(api_config or {})
        super().__init__(merged_config)
        self.ai_case = ai_case or {}

    def validate_config(self):
        return self.ai_case.get("validate_ok", True), self.ai_case.get("validate_message", "ok")

    def send_request(self, prompt):
        if not self.ai_case.get("request_ok", True):
            return False, self.ai_case.get("request_error", "replayed request failure")

        response = self.ai_case.get("response")
        if response is None:
            response = {
                "choices": [
                    {
                        "message": {
                            "content": self.ai_case.get("response_content", ""),
                        }
                    }
                ]
            }
        return True, response


class EvalDocProcessor(DocProcessor):
    """Doc processor that captures rendered output details for eval assertions."""

    ALIGNMENT_VALUES = {
        0: "left",
        1: "center",
        2: "right",
        3: "justify",
    }

    def apply_formatting(self, formatting_instructions, custom_save_path=None, header_footer_config=None):
        report = super().apply_formatting(
            formatting_instructions,
            custom_save_path=custom_save_path,
            header_footer_config=header_footer_config,
        )

        output_file = report.get("output_file")
        if report.get("success") and output_file and Path(output_file).exists():
            report["captured_output"] = self._capture_output(output_file)

        return report

    def _capture_output(self, output_file):
        document = Document(output_file)
        paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
        return {
            "paragraph_texts": [paragraph.text for paragraph in paragraphs],
            "alignments": [self._normalize_alignment(paragraph.alignment) for paragraph in paragraphs],
        }

    def _normalize_alignment(self, alignment):
        if alignment is None:
            return "left"
        value = int(alignment)
        return self.ALIGNMENT_VALUES.get(value, "left")


class RuntimeEvalHarness:
    """Run replayable eval cases through the runtime harness."""

    def __init__(self, runtime_dir="runtime", eval_dir="runtime/evals"):
        self.runtime_dir = Path(runtime_dir)
        self.eval_dir = Path(eval_dir)

    def load_cases(self, cases_dir):
        cases_path = Path(cases_dir)
        cases = []
        for case_path in sorted(cases_path.glob("*.json")):
            case = json.loads(case_path.read_text(encoding="utf-8"))
            if case.get("schema_version") != "formulaai.eval_case.v1":
                raise ValueError(f"unsupported eval case schema: {case_path}")
            case["_path"] = str(case_path)
            cases.append(case)
        return cases

    def run_suite(self, cases_dir):
        cases = self.load_cases(cases_dir)
        suite_id = f"{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_{uuid4().hex[:8]}"
        summary_dir = self.eval_dir / suite_id
        summary_dir.mkdir(parents=True, exist_ok=True)

        results = [self._run_case(case) for case in cases]
        summary = {
            "schema_version": "formulaai.eval_suite.v1",
            "suite_id": suite_id,
            "created_at": _iso_now(),
            "cases_total": len(results),
            "cases_passed": sum(1 for result in results if result["passed"]),
            "cases_failed": sum(1 for result in results if not result["passed"]),
            "results": results,
        }

        summary_path = summary_dir / "summary.json"
        results_path = summary_dir / "results.jsonl"
        summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
        with results_path.open("w", encoding="utf-8") as handle:
            for result in results:
                handle.write(json.dumps(result, ensure_ascii=False) + "\n")
        summary["summary_path"] = str(summary_path)
        summary["results_path"] = str(results_path)
        return summary

    def _run_case(self, case):
        source_bytes = self._build_source_docx(case["source"]["paragraphs"])
        ai_case = case.get("ai", {})
        api_config = {
            "api_url": "https://eval.local/v1/chat/completions",
            "api_key": "eval-key",
            "model": "eval-model",
            "timeout": 1,
        }
        harness = DocumentFormatHarness(
            runtime_dir=self.runtime_dir,
            doc_processor_factory=EvalDocProcessor,
            ai_connector_factory=lambda config, replay=ai_case: ReplayAIConnector(config, replay),
        )

        result = harness.run(
            source_name=case["source"]["source_name"],
            source_bytes=source_bytes,
            template_name=case["template"]["name"],
            template_rules=case["template"]["rules"],
            api_config=api_config,
            header_footer_config=case.get("header_footer_config", {}),
        )

        manifest_path = self._find_manifest_path(result.run_id)
        manifest = self._read_json(manifest_path) if manifest_path else {}
        actual = self._build_actual(case, result, manifest, api_config)
        mismatches = self._compare_expected(case.get("expect", {}), actual)

        return {
            "case_id": case["id"],
            "description": case.get("description", ""),
            "passed": not mismatches,
            "mismatches": mismatches,
            "actual": actual,
            "run_id": result.run_id,
            "manifest_path": str(manifest_path) if manifest_path else None,
            "case_path": case.get("_path"),
        }

    def _build_source_docx(self, paragraphs):
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _find_manifest_path(self, run_id):
        matches = sorted(self.runtime_dir.rglob(f"{run_id}/manifest.json"))
        return matches[-1] if matches else None

    def _build_actual(self, case, result, manifest, api_config):
        output_validation = manifest.get("output_validation", {})
        render_report = result.render_report or {}
        captured_output = render_report.get("captured_output", {})

        return {
            "status": _enum_value(result.status),
            "final_stage": _enum_value(result.final_stage),
            "instruction_count": result.instruction_count,
            "processed_elements": render_report.get("processed_elements", 0),
            "error_code": _enum_value(result.error_code),
            "output_document_loadable": output_validation.get("document_loadable"),
            "output_paragraph_count": output_validation.get("paragraph_count"),
            "output_alignments": captured_output.get("alignments") or self._extract_output_alignments(case, api_config),
            "output_paragraph_texts": captured_output.get("paragraph_texts", []),
        }

    def _extract_output_alignments(self, case, api_config):
        connector = ReplayAIConnector(api_config, case.get("ai", {}))
        success, payload = connector.send_request("")
        if not success:
            return []

        success, formatting_instructions = connector.parse_response(payload)
        if not success:
            return []

        return [
            element.get("format", {}).get("alignment")
            for element in formatting_instructions.get("elements", [])
            if element.get("format", {}).get("alignment") is not None
        ]

    def _compare_expected(self, expected, actual):
        mismatches = []
        for field, expected_value in expected.items():
            actual_value = actual.get(field)
            if actual_value != expected_value:
                mismatches.append(
                    {
                        "field": field,
                        "expected": expected_value,
                        "actual": actual_value,
                    }
                )
        return mismatches

    def _read_json(self, path):
        return json.loads(Path(path).read_text(encoding="utf-8"))


def main(argv=None):
    parser = argparse.ArgumentParser(description="Run offline FormulaAI eval cases.")
    parser.add_argument("--cases-dir", default="evals/cases")
    parser.add_argument("--runtime-dir", default="runtime")
    parser.add_argument("--eval-dir", default="runtime/evals")
    args = parser.parse_args(argv)

    suite_result = RuntimeEvalHarness(runtime_dir=args.runtime_dir, eval_dir=args.eval_dir).run_suite(args.cases_dir)
    print(json.dumps({"summary_path": suite_result["summary_path"], "cases_failed": suite_result["cases_failed"]}, ensure_ascii=False))
    return 0 if suite_result["cases_failed"] == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
